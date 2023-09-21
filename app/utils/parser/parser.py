import os

from app.utils.document.document_utils import contains_only_punctuation
from app.utils.file.file_utils import allowed_file, create_path
from docx import Document
from PyPDF2 import PdfReader

import pandas as pd
import docx
import re


def extract_file_content(file_path, file_name):
    # Data la path e il nome di un file ritorna il testo contenuto.

    if allowed_file(file_name, {"xls", 'xlsx'}):
        df = pd.read_excel(file_path)
        file_text = df.to_string()
    elif allowed_file(file_name, {"csv"}):
        df = pd.read_csv(file_path)
        file_text = df.to_string()
    elif allowed_file(file_name, {"docx"}):
        doc = docx.Document(file_path)
        file_text = ""
        for paragraph in doc.paragraphs:
            file_text += paragraph.text + "\n"
    elif allowed_file(file_name, {"pdf"}):
        file_text = extract_text_from_pdf(file_path)
    else:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                file_text = f.read()
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='latin-1') as f:
                file_text = f.read()
    return file_text


def extract_text_from_pdf(pdf_path):
    # Data la path di un pdf ritorna il contenuto del pdf

    with open(pdf_path, 'rb') as file:
        reader = PdfReader(file)
        text = ""
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            text += page.extract_text()

    return text


def extract_text_from_docx(docx_file):
    # Dato il contenuto di un Documento ritorna una lista contenente run e paragrafi
    # Include tabelle, Footer , Header e Testo
    try:
        parts = []
        doc = Document(docx_file)

        for paragraph in doc.paragraphs:
            if paragraph.text and re.sub(r'[ \t]', '', paragraph.text) != "" and not (contains_only_punctuation(paragraph.text.replace(" ", ""))):
                parts.append(paragraph.text)

        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                if paragraph.text and re.sub(r'[ \t]', '', paragraph.text) != "" and not (
                            contains_only_punctuation(paragraph.text.replace(" ", ""))):
                    parts.append(paragraph.text )

            footer = section.footer
            for paragraph in footer.paragraphs:
                if paragraph.text and re.sub(r'[ \t]', '', paragraph.text) != "" and not (
                        contains_only_punctuation(paragraph.text.replace(" ", ""))):
                    parts.append(paragraph.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if paragraph.text and re.sub(r'[ \t]', '', paragraph.text) != "" and not (
                                    contains_only_punctuation(paragraph.text.replace(" ", ""))):
                                parts.append(paragraph.text)
        return parts
    except Exception as e:
        raise e


def write_text_to_docx(docx_file, translations, type):
    # Dato il file e una lista con le traduzioni, modifica ogni run e paragrafo del file con la corrispettiva traduzione.

    doc = Document(docx_file)
    translated_index = 0

    for paragraph in doc.paragraphs:
        if paragraph.text and re.sub(r'[ \t]', '',  paragraph.text) != "" and not contains_only_punctuation(
                paragraph.text.replace(" ", "")):

            translation = translations[translated_index]
            translated_index += 1
            if paragraph.text.startswith(' ') and not translation.startswith(' '):
                translation = ' ' + translation
            if paragraph.text.endswith(' ') and not translation.endswith(' '):
                translation += ' '
            paragraph.text = translation

    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            if paragraph.text and re.sub(r'[ \t]', '', paragraph.text) != "" and not (
                    contains_only_punctuation(paragraph.text.replace(" ", ""))):
                    translation = translations[translated_index]
                    translated_index += 1

                    if paragraph.text.startswith(' ') and not translation.startswith(' '):
                        translation = ' ' + translation
                    if paragraph.text.endswith(' ') and not translation.endswith(' '):
                        translation += ' '
                    paragraph.text = translation

        footer = section.footer
        for paragraph in footer.paragraphs:
            if paragraph.text and re.sub(r'[ \t]', '', paragraph.text) != "" and not (
                    contains_only_punctuation(paragraph.text.replace(" ", ""))):
                translation = translations[translated_index]
                translated_index += 1

                if paragraph.text.startswith(' ') and not translation.startswith(' '):
                    translation = ' ' + translation
                if paragraph.text.endswith(' ') and not translation.endswith(' '):
                    translation += ' '
                paragraph.text = translation

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text and re.sub(r'[ \t]', '', paragraph.text) != "" and not (
                            contains_only_punctuation(paragraph.text.replace(" ", ""))):
                            translation = translations[translated_index]
                            translated_index += 1

                            if paragraph.text.startswith(' ') and not translation.startswith(' '):
                                translation = ' ' + translation
                            if paragraph.text.endswith(' ') and not translation.endswith(' '):
                                translation += ' '
                            paragraph.text = translation

    file_path = create_path(".docx", type)
    relative_path = os.path.relpath(file_path, start="myprog")
    if not os.path.exists(os.path.dirname(file_path)):
        os.makedirs(os.path.dirname(file_path))
    doc.save(file_path)
    return relative_path